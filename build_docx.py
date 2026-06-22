from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "文档.md"
OUTPUT = ROOT / "文档.docx"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name="Microsoft YaHei", size=None, color=None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = RGBColor(*color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_inline_markdown(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_east_asia_font(run, "Consolas")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(68, 68, 68)
        else:
            run = paragraph.add_run(part)
            set_east_asia_font(run)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"图：{text}")
    set_east_asia_font(run)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def image_width(path):
    with Image.open(path) as img:
        w, h = img.size
    aspect = w / h
    max_width = 6.15
    max_height = 6.55
    return Inches(min(max_width, max_height * aspect))


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    set_style_font(styles["Normal"], size=11, color=(34, 34, 34))
    styles["Normal"].paragraph_format.line_spacing = 1.12
    styles["Normal"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Heading 1"], size=16, color=(46, 116, 181))
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)

    set_style_font(styles["Heading 2"], size=13, color=(31, 77, 120))
    styles["Heading 2"].paragraph_format.space_before = Pt(11)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    set_style_font(styles["Heading 3"], size=12, color=(31, 77, 120))
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ["List Bullet", "List Number"]:
        set_style_font(styles[style_name], size=11, color=(34, 34, 34))
        styles[style_name].paragraph_format.space_after = Pt(4)
        styles[style_name].paragraph_format.line_spacing = 1.12

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    pending_para = []
    first_heading_done = False

    def flush_para():
        nonlocal pending_para
        if not pending_para:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        add_inline_markdown(p, "".join(pending_para).strip())
        pending_para = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_para()
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            flush_para()
            alt, rel = image_match.groups()
            img_path = (ROOT / rel.replace("./", "")).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            run.add_picture(str(img_path), width=image_width(img_path))
            add_caption(doc, alt)
            continue

        if line.startswith("# "):
            flush_para()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[2:].strip())
            set_east_asia_font(run)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(31, 77, 120)
            first_heading_done = True
            continue

        if line.startswith("## "):
            flush_para()
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue

        if line.startswith("### "):
            flush_para()
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if line.startswith("> "):
            flush_para()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(9)
            run = p.add_run(line[2:].strip())
            set_east_asia_font(run)
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
            continue

        bullet_match = re.match(r"^-\s+(.*)", line)
        if bullet_match:
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet_match.group(1))
            continue

        number_match = re.match(r"^\d+\.\s+(.*)", line)
        if number_match:
            flush_para()
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, number_match.group(1))
            continue

        if pending_para:
            pending_para.append(line.strip())
        else:
            pending_para.append(line.strip())

    flush_para()
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
